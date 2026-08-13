"""Deterministic document corpus for the AI workflow (prompt-caching aware).

Every step of the tender workflow needs to read the same project documents.
Sending them to the model as raw files on every call would be slow and
expensive, so instead we convert every document to plain text ONCE per run
and concatenate the results into a single "corpus" string.

The corpus is the biggest cached block in every model request, and Bedrock
prompt caching only hits when the bytes are IDENTICAL between requests, so
determinism is the whole game here:

  * documents are processed in sorted-filename order (never dict/upload order);
  * formatting is fixed — no timestamps, job ids or randomness anywhere;
  * the finished corpus is persisted in the project's storage folder under
    ``CORPUS_FILENAME`` with a content-hash header. On the next run we re-use
    the stored bytes verbatim when the hash still matches, so even a library
    upgrade that would extract slightly different text cannot silently change
    the corpus mid-way through a set of runs.

Generated documents (the compiled "Project Context …" file) are excluded —
they contain a generation timestamp and their content is already passed to
the model separately, so including them would bust the cache on every run.
"""

from __future__ import annotations

import hashlib
import io
import logging
from pathlib import Path

from app import settings

LOGGER = logging.getLogger(__name__)

CORPUS_FILENAME = "_document_corpus.txt"

# Bump when extraction logic changes so stale stored corpora are rebuilt.
EXTRACTOR_VERSION = "1"

_XLSX_EXTS = {".xlsx", ".xlsm", ".xltx"}
_TEXT_EXTS = {".txt", ".csv", ".md", ".rtf"}
_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".heic"}


# --------------------------------------------------------------------------- #
# Per-format text extraction                                                  #
# --------------------------------------------------------------------------- #

def extract_text(filename: str, data: bytes) -> str:
    """Best-effort plain-text extraction for one document. Deterministic for
    the same (filename, bytes). Never raises."""
    ext = Path(filename).suffix.lower()
    try:
        if ext == ".pdf":
            return _from_pdf(data)
        if ext in (".docx", ".doc"):
            return _from_docx(data) if ext == ".docx" else (
                "[.doc (legacy Word) file — convert to .docx for text extraction]")
        if ext in _XLSX_EXTS or ext == ".xls":
            return _from_xlsx(data) if ext in _XLSX_EXTS else (
                "[.xls (legacy Excel) file — convert to .xlsx for text extraction]")
        if ext in _TEXT_EXTS:
            return data.decode("utf-8", errors="replace")
        if ext in _IMAGE_EXTS:
            return ("[image file — no text extracted; OCR is not enabled. "
                    "Note the filename and treat contents as unavailable]")
        return f"[{ext or 'unknown'} file — not text-extractable]"
    except Exception as exc:  # noqa: BLE001 — a bad file must not kill the run
        LOGGER.warning("Text extraction failed for %s: %s", filename, exc)
        return f"[extraction failed: {exc.__class__.__name__}]"


def _from_pdf(data: bytes) -> str:
    import pdfplumber  # lazy import

    pages: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = (page.extract_text() or "").strip()
            pages.append(f"--- page {i} ---\n{text}" if text
                         else f"--- page {i} ---\n[no extractable text — scanned page?]")
    return "\n\n".join(pages) if pages else "[empty PDF]"


def _from_docx(data: bytes) -> str:
    import docx  # python-docx, lazy import

    d = docx.Document(io.BytesIO(data))
    parts: list[str] = [p.text for p in d.paragraphs if p.text.strip()]
    for t_i, table in enumerate(d.tables, start=1):
        parts.append(f"--- table {t_i} ---")
        for row in table.rows:
            parts.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(parts) or "[empty document]"


def _from_xlsx(data: bytes, max_rows_per_sheet: int = 400) -> str:
    from openpyxl import load_workbook  # lazy import

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    try:
        for ws in wb.worksheets:
            parts.append(f"--- sheet: {ws.title} ---")
            for r_i, row in enumerate(ws.iter_rows(values_only=True)):
                if r_i >= max_rows_per_sheet:
                    parts.append(f"[…truncated at {max_rows_per_sheet} rows]")
                    break
                cells = ["" if v is None else str(v) for v in row]
                if any(c.strip() for c in cells):
                    # Cell reference prefix helps the model cite "row 42".
                    parts.append(f"r{r_i + 1}: " + " | ".join(cells).rstrip(" |"))
    finally:
        wb.close()
    return "\n".join(parts) or "[empty workbook]"


# --------------------------------------------------------------------------- #
# Corpus assembly                                                             #
# --------------------------------------------------------------------------- #

def _corpus_hash(docs: list[dict], char_cap: int) -> str:
    """Content hash over the inputs that determine the corpus bytes."""
    h = hashlib.sha256()
    h.update(f"v{EXTRACTOR_VERSION}|cap{char_cap}".encode())
    for d in sorted(docs, key=lambda d: d["filename"]):
        h.update(d["filename"].encode("utf-8", errors="replace"))
        h.update(hashlib.sha256(d["bytes"]).digest())
    return h.hexdigest()[:32]


def _build(docs: list[dict], char_cap: int) -> str:
    parts: list[str] = []
    for d in sorted(docs, key=lambda d: d["filename"]):
        text = extract_text(d["filename"], d["bytes"])
        if len(text) > char_cap:
            text = (text[:char_cap]
                    + f"\n[…document truncated at {char_cap} characters — "
                      "use the read_document tool for the full text]")
        parts.append(f"===== FILE: {d['filename']} =====\n{text}")
    return "\n\n".join(parts) if parts else "[no project documents uploaded]"


def get_or_build(pid: str, documents: list[dict],
                 generated_names: set[str] | None = None) -> str:
    """Return the corpus for a project, re-using the stored copy when the
    document set is unchanged (byte-identical ⇒ cache hits across runs).

    `documents` is the [{filename, bytes}] payload list handed to the AI;
    `generated_names` are filenames of app-generated docs to exclude.
    """
    from app.features.projects import repo as _repo  # lazy — avoid cycles

    generated_names = generated_names or set()
    docs = [d for d in documents
            if d["filename"] not in generated_names
            and d["filename"] != CORPUS_FILENAME]

    char_cap = settings.AI_CORPUS_DOC_CHAR_CAP
    want_hash = _corpus_hash(docs, char_cap)
    header = f"# corpus {want_hash}\n"

    repo = _repo.get_repo()
    stored = repo.read_document(pid, CORPUS_FILENAME)
    if stored is not None:
        text = stored.decode("utf-8", errors="replace")
        if text.startswith(header):
            LOGGER.info("Corpus re-used for %s (hash %s).", pid, want_hash)
            return text

    body = header + _build(docs, char_cap)
    repo.write_document(pid, CORPUS_FILENAME, body.encode("utf-8"))
    LOGGER.info("Corpus built for %s: %d docs, %d chars (hash %s).",
                pid, len(docs), len(body), want_hash)
    return body


def full_document_text(pid: str, filename: str) -> str:
    """Uncapped extraction of one document — backs the read_document tool."""
    from app.features.projects import repo as _repo  # lazy — avoid cycles

    data = _repo.get_repo().read_document(pid, filename)
    if data is None:
        return f"[document not found: {filename}]"
    return extract_text(filename, data)
